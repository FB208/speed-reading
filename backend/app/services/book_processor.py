import re
import os
import hashlib
import shutil
import posixpath
from urllib.parse import unquote, urlparse
from typing import List
from docx import Document
from sqlalchemy.orm import Session
from app.models import models


class BookProcessor:
    """书籍处理器：提取正文并分段（保留格式）"""

    def __init__(self, db: Session):
        self.db = db
        self.uploads_dir = os.path.abspath(
            os.path.join(os.path.dirname(__file__), "..", "..", "uploads")
        )
        self.book_images_root = os.path.join(self.uploads_dir, "book_images")
        os.makedirs(self.book_images_root, exist_ok=True)
        self._saved_image_map = {}

    async def process_book(self, book_id: int, file_path: str):
        """处理书籍文件"""
        self._saved_image_map = {}
        # 读取文件内容（保留HTML格式）
        content = self._read_file(file_path, book_id)

        # 不再清理内容
        cleaned_content = content

        # 分段（保留HTML标签）
        paragraphs = self._split_into_paragraphs(cleaned_content)

        # 保存段落到数据库
        for i, paragraph_content in enumerate(paragraphs, 1):
            # 计算纯文本字数（去除HTML标签）
            plain_text = self._strip_html_tags(paragraph_content)
            word_count = len(plain_text)

            paragraph = models.Paragraph(
                book_id=book_id,
                sequence=i,
                content=paragraph_content,
                word_count=word_count,
            )
            self.db.add(paragraph)

        # 更新书籍段落数
        book = self.db.query(models.Book).filter(models.Book.id == book_id).first()
        if book:
            book.total_paragraphs = len(paragraphs)

        self.db.commit()

    def _strip_html_tags(self, html_content: str) -> str:
        """去除HTML标签，只保留纯文本"""
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(html_content, "html.parser")
            return soup.get_text()
        except:
            # 如果BeautifulSoup不可用，使用简单的正则
            clean = re.sub(r"<[^>]+>", "", html_content)
            return clean

    def _read_file(self, file_path: str, book_id: int) -> str:
        """读取文件内容"""
        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext == ".txt":
            # 尝试多种编码
            encodings = ["utf-8", "gbk", "gb2312", "utf-16"]
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        content = f.read()
                        # 将markdown格式转换为HTML
                        return self._convert_markdown_to_html(content)
                except UnicodeDecodeError:
                    continue
            raise Exception("无法识别文件编码")

        elif file_ext == ".docx":
            return self._read_docx_with_format(file_path)

        elif file_ext == ".epub":
            return self._read_epub_with_format(file_path, book_id)

        elif file_ext == ".mobi":
            return self._read_mobi_with_format(file_path, book_id)

        elif file_ext == ".pdf":
            return self._read_pdf(file_path)

        else:
            raise Exception(f"不支持的文件格式: {file_ext}")

    def _convert_markdown_to_html(self, content: str) -> str:
        """将Markdown格式转换为HTML"""
        # 转换标题
        content = re.sub(
            r"^#{1,6}\s+(.+)$",
            lambda m: f"<h{len(m.group(0)) - len(m.group(1)) - 1}>{m.group(1)}</h{len(m.group(0)) - len(m.group(1)) - 1}>",
            content,
            flags=re.MULTILINE,
        )

        # 转换加粗 **text** 或 __text__
        content = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", content)
        content = re.sub(r"__(.+?)__", r"<strong>\1</strong>", content)

        # 转换斜体 *text* 或 _text_
        content = re.sub(r"\*(.+?)\*", r"<em>\1</em>", content)
        content = re.sub(r"_(.+?)_", r"<em>\1</em>", content)

        # 转换段落（将空行分隔的文本包装为<p>标签）
        paragraphs = content.split("\n\n")
        result = []
        for para in paragraphs:
            para = para.strip()
            if para and not para.startswith("<"):
                para = f"<p>{para}</p>"
            result.append(para)

        return "\n\n".join(result)

    def _read_docx_with_format(self, file_path: str) -> str:
        """读取Word文档并保留格式"""
        try:
            doc = Document(file_path)
            html_parts = []

            for para in doc.paragraphs:
                if not para.text.strip():
                    continue

                # 检测标题样式
                style_name = para.style.name.lower()
                text = para.text

                # 应用格式
                formatted_text = text
                for run in para.runs:
                    run_text = run.text
                    if run.bold:
                        run_text = f"<strong>{run_text}</strong>"
                    if run.italic:
                        run_text = f"<em>{run_text}</em>"
                    # 替换原文中的这段文本
                    formatted_text = formatted_text.replace(run.text, run_text)

                # 根据样式添加标签
                if "heading 1" in style_name or "标题 1" in style_name:
                    html_parts.append(f"<h1>{formatted_text}</h1>")
                elif "heading 2" in style_name or "标题 2" in style_name:
                    html_parts.append(f"<h2>{formatted_text}</h2>")
                elif "heading 3" in style_name or "标题 3" in style_name:
                    html_parts.append(f"<h3>{formatted_text}</h3>")
                elif "heading" in style_name or "标题" in style_name:
                    html_parts.append(f"<h3>{formatted_text}</h3>")
                else:
                    html_parts.append(f"<p>{formatted_text}</p>")

            return "\n".join(html_parts)
        except Exception as e:
            # 如果失败，返回纯文本
            doc = Document(file_path)
            return "\n".join(
                [f"<p>{p.text}</p>" for p in doc.paragraphs if p.text.strip()]
            )

    def _read_epub_with_format(self, file_path: str, book_id: int) -> str:
        """读取EPUB文件并保留HTML格式（包含图片）"""
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup

            book = epub.read_epub(file_path)
            html_parts = []
            image_items = {}

            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_IMAGE:
                    item_name = self._normalize_epub_path(item.get_name())
                    image_items[item_name] = item

            for item in book.get_items():
                if item.get_type() != ebooklib.ITEM_DOCUMENT:
                    continue

                content = item.get_content().decode("utf-8", errors="ignore")
                soup = BeautifulSoup(content, "html.parser")

                for tag in soup(["script", "style", "svg"]):
                    tag.decompose()

                doc_name = self._normalize_epub_path(item.get_name())
                for img in soup.find_all("img"):
                    src = (img.get("src") or "").strip()
                    if not src or self._is_external_image(src):
                        continue

                    image_item = self._resolve_epub_image_item(
                        src, doc_name, image_items
                    )
                    if not image_item:
                        continue

                    image_url = self._save_image_bytes(
                        book_id=book_id,
                        source_name=image_item.get_name(),
                        image_data=image_item.get_content(),
                    )
                    if image_url:
                        img["src"] = image_url
                        if img.has_attr("srcset"):
                            del img["srcset"]

                body = soup.find("body")
                html_parts.append(str(body) if body else str(soup))

            return "\n".join(html_parts)
        except Exception as e:
            raise Exception(f"读取EPUB文件失败: {str(e)}")

    def _read_mobi_with_format(self, file_path: str, book_id: int) -> str:
        """读取MOBI文件并保留HTML格式"""
        try:
            import mobi
            from bs4 import BeautifulSoup

            # 使用mobi库提取内容
            tempdir, filepath = mobi.extract(file_path)

            try:
                with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()

                soup = BeautifulSoup(content, "html.parser")

                for tag in soup(["script", "style", "svg"]):
                    tag.decompose()

                html_dir = os.path.dirname(filepath)
                tempdir_abs = os.path.abspath(tempdir)
                for img in soup.find_all("img"):
                    src = (img.get("src") or "").strip()
                    if not src or self._is_external_image(src):
                        continue

                    parsed = urlparse(src)
                    relative_path = unquote(parsed.path).lstrip("/")
                    if not relative_path:
                        continue

                    image_path = os.path.normpath(os.path.join(html_dir, relative_path))
                    if not os.path.exists(image_path):
                        image_path = os.path.normpath(
                            os.path.join(tempdir, relative_path)
                        )
                    image_path_abs = os.path.abspath(image_path)
                    if os.path.commonpath([tempdir_abs, image_path_abs]) != tempdir_abs:
                        continue
                    if not os.path.isfile(image_path_abs):
                        continue

                    with open(image_path_abs, "rb") as image_file:
                        image_data = image_file.read()

                    image_url = self._save_image_bytes(
                        book_id=book_id,
                        source_name=os.path.basename(image_path_abs),
                        image_data=image_data,
                    )
                    if image_url:
                        img["src"] = image_url
                        if img.has_attr("srcset"):
                            del img["srcset"]

                return str(soup)
            finally:
                shutil.rmtree(tempdir, ignore_errors=True)
        except Exception as e:
            raise Exception(f"读取MOBI文件失败: {str(e)}")

    def _normalize_epub_path(self, path: str) -> str:
        return posixpath.normpath((path or "").replace("\\", "/")).lstrip("./")

    def _resolve_epub_image_item(self, src: str, doc_name: str, image_items: dict):
        src_path = self._normalize_epub_path(urlparse(src).path)
        if not src_path:
            return None

        doc_dir = posixpath.dirname(doc_name)
        candidates = [
            self._normalize_epub_path(posixpath.join(doc_dir, src_path)),
            src_path,
        ]

        for candidate in candidates:
            if candidate in image_items:
                return image_items[candidate]

        basename = posixpath.basename(src_path)
        for name, item in image_items.items():
            if posixpath.basename(name) == basename:
                return item
        return None

    def _is_external_image(self, src: str) -> bool:
        normalized = src.lower().strip()
        return normalized.startswith(("http://", "https://", "data:", "/book-images/"))

    def _guess_image_ext(self, source_name: str, image_data: bytes) -> str:
        ext = os.path.splitext(source_name or "")[1].lower()
        if ext in [".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".svg"]:
            return ext
        if image_data.startswith(b"\x89PNG\r\n\x1a\n"):
            return ".png"
        if image_data.startswith(b"\xff\xd8"):
            return ".jpg"
        if image_data.startswith(b"GIF8"):
            return ".gif"
        if image_data.startswith(b"RIFF"):
            return ".webp"
        if image_data.startswith(b"BM"):
            return ".bmp"
        return ".jpg"

    def _save_image_bytes(
        self, book_id: int, source_name: str, image_data: bytes
    ) -> str:
        if not image_data:
            return ""

        digest = hashlib.md5(image_data).hexdigest()
        cache_key = f"{book_id}:{digest}"
        if cache_key in self._saved_image_map:
            return self._saved_image_map[cache_key]

        ext = self._guess_image_ext(source_name, image_data)
        filename = f"img_{digest[:16]}{ext}"
        book_dir_name = f"book_{book_id}"
        book_dir = os.path.join(self.book_images_root, book_dir_name)
        os.makedirs(book_dir, exist_ok=True)
        file_path = os.path.join(book_dir, filename)

        if not os.path.exists(file_path):
            with open(file_path, "wb") as f:
                f.write(image_data)

        url = f"/book-images/{book_dir_name}/{filename}"
        self._saved_image_map[cache_key] = url
        return url

    def _read_pdf(self, file_path: str) -> str:
        """读取PDF文件（PDF不支持富文本格式）"""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            html_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    # 将PDF文本包装为HTML段落
                    paragraphs = text.split("\n\n")
                    for para in paragraphs:
                        if para.strip():
                            html_parts.append(f"<p>{para.strip()}</p>")

            return "\n".join(html_parts)
        except Exception as e:
            raise Exception(f"读取PDF文件失败: {str(e)}")

    def _clean_content(self, content: str) -> str:
        """清理内容：保留全部内容，不进行自动清理"""
        return content

    def _split_into_paragraphs(self, content: str) -> List[str]:
        """
        将内容分段（保留HTML格式）
        规则：按自然段拆分，如果纯文本不足1000字则合并下一个自然段
        """
        # 对于HTML内容，按<p>, <div>, <h1>-<h6>等标签分割
        if "<" in content and ">" in content:
            return self._split_html_paragraphs(content)
        else:
            # 纯文本分段
            return self._split_text_paragraphs(content)

    def _split_html_paragraphs(self, content: str) -> List[str]:
        """分割HTML段落（保留标签）"""
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(content, "html.parser")

            # 获取所有顶级元素
            elements = []
            for elem in soup.find_all(
                [
                    "p",
                    "div",
                    "h1",
                    "h2",
                    "h3",
                    "h4",
                    "h5",
                    "h6",
                    "li",
                    "blockquote",
                    "figure",
                    "img",
                ]
            ):
                if elem.parent.name in [
                    "body",
                    "html",
                    "[document]",
                ] or elem.parent.name not in ["p", "div"]:
                    elements.append(str(elem))

            if not elements:
                # 如果没有找到合适元素，按文本回退分段，避免单段过长
                return self._split_oversized_paragraphs([content])

            # 合并段落（根据纯文本字数）
            result = []
            current_paragraph = ""
            current_text_length = 0

            for elem_str in elements:
                elem_soup = BeautifulSoup(elem_str, "html.parser")
                text_length = len(elem_soup.get_text())

                if not current_paragraph:
                    current_paragraph = elem_str
                    current_text_length = text_length
                else:
                    # 如果当前段落不足1000字，合并
                    if current_text_length < 1000:
                        current_paragraph += "\n" + elem_str
                        current_text_length += text_length
                    else:
                        # 当前段落已足够长，保存并开始新段落
                        result.append(current_paragraph)
                        current_paragraph = elem_str
                        current_text_length = text_length

            # 处理最后一个段落
            if current_paragraph:
                result.append(current_paragraph)

            return self._split_oversized_paragraphs(result if result else [content])
        except Exception as e:
            # 如果解析失败，按文本回退分段，避免单段过长
            return self._split_oversized_paragraphs([content])

    def _split_oversized_paragraphs(
        self, paragraphs: List[str], max_text_length: int = 4000
    ) -> List[str]:
        """将超长段落按纯文本长度切分，避免数据库字段超限"""
        try:
            from bs4 import BeautifulSoup

            normalized: List[str] = []
            for paragraph in paragraphs:
                text = BeautifulSoup(paragraph, "html.parser").get_text("\n")
                text_blocks = [
                    line.strip() for line in text.split("\n") if line.strip()
                ]

                if not text_blocks:
                    continue

                current = ""
                for block in text_blocks:
                    if not current:
                        current = block
                        continue

                    if len(current) + len(block) + 1 <= max_text_length:
                        current += "\n" + block
                    else:
                        normalized.append(f"<p>{current}</p>")
                        current = block

                if current:
                    normalized.append(f"<p>{current}</p>")

            return normalized if normalized else paragraphs
        except Exception:
            # 回退到最基础按字符切分
            normalized = []
            for paragraph in paragraphs:
                if len(paragraph) <= max_text_length:
                    normalized.append(paragraph)
                    continue

                for i in range(0, len(paragraph), max_text_length):
                    chunk = paragraph[i : i + max_text_length]
                    normalized.append(f"<p>{chunk}</p>")

            return normalized if normalized else paragraphs

    def _split_text_paragraphs(self, content: str) -> List[str]:
        """分割纯文本段落"""
        # 先按自然段分割
        raw_paragraphs = [p.strip() for p in content.split("\n") if p.strip()]

        result = []
        current_paragraph = ""

        for raw_para in raw_paragraphs:
            # 如果当前段落为空，直接添加
            if not current_paragraph:
                current_paragraph = raw_para
            else:
                # 如果当前段落不足1000字，合并
                if len(current_paragraph) < 1000:
                    current_paragraph += "\n" + raw_para
                else:
                    # 当前段落已足够长，保存并开始新段落
                    result.append(current_paragraph)
                    current_paragraph = raw_para

        # 处理最后一个段落
        if current_paragraph:
            result.append(current_paragraph)

        return result
